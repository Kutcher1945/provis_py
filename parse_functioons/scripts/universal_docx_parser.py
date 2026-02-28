#!/usr/bin/env python3
"""
Universal DOCX parser that handles all document structure types

Supports:
- CHAPTER-BASED: "Глава 1", "Глава 2"...
- NUMBERED SECTIONS: "1. Общие положения", "2. Задачи"...
- CUSTOM: Keyword-based extraction
"""

import sys
from docx import Document
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))


def detect_structure_type(doc):
    """Detect document structure type"""

    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)

    full_text = '\n'.join(all_text).lower()

    # Detect structure
    if 'глава 1' in full_text and 'глава 2' in full_text:
        return 'CHAPTER-BASED'
    elif '1. общие положения' in full_text or '1.общие положения' in full_text:
        return 'NUMBERED-SECTIONS'
    else:
        return 'CUSTOM'


def parse_chapter_based(doc):
    """Parse CHAPTER-BASED documents (current УАиГ logic)"""

    general_provisions = []
    tasks = []
    authorities_rights = []
    authorities_responsibilities = []
    functions = []
    additions = []

    current_chapter = None
    current_subsection = None
    in_rights = False
    in_responsibilities = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect chapters
        if 'Глава 1' in text or 'ГЛАВА 1' in text:
            current_chapter = 1
            current_subsection = None
            continue
        elif 'Глава 2' in text or 'ГЛАВА 2' in text:
            current_chapter = 2
            current_subsection = None
            continue
        elif 'Глава 3' in text or 'ГЛАВА 3' in text:
            current_chapter = 3
            current_subsection = None
            additions.append(text)
            continue
        elif 'Глава 4' in text or 'ГЛАВА 4' in text:
            current_chapter = 4
            current_subsection = None
            additions.append(text)
            continue
        elif 'Глава 5' in text or 'ГЛАВА 5' in text:
            current_chapter = 5
            current_subsection = None
            additions.append(text)
            continue

        # Detect subsections in Глава 2
        if current_chapter == 2:
            if 'Задачи' in text and (text.startswith('12') or text.startswith('Задачи')):
                current_subsection = 'tasks'
                continue
            elif 'Полномочия' in text and (text.startswith('13') or text.startswith('Полномочия')):
                current_subsection = 'authorities'
                in_rights = False
                in_responsibilities = False
                # Check if "права:" is on the same line
                if 'права:' in text.lower() or 'права :' in text.lower():
                    in_rights = True
                continue
            elif 'Функции' in text and (text.startswith('14') or text.startswith('15') or text.startswith('Функции')):
                current_subsection = 'functions'
                in_rights = False
                in_responsibilities = False
                continue

        # Extract based on current location
        if current_chapter == 1:
            general_provisions.append(text)

        elif current_chapter == 2:
            if current_subsection == 'tasks':
                if text and text[0].isdigit() and ')' in text[:5]:
                    tasks.append(text)

            elif current_subsection == 'authorities':
                if 'права:' in text.lower() or 'права :' in text.lower():
                    in_rights = True
                    in_responsibilities = False
                    continue
                elif 'обязанности:' in text.lower() or 'обязанности :' in text.lower():
                    in_rights = False
                    in_responsibilities = True
                    continue

                # Collect both numbered (1., 2., 3.) and unnumbered items
                if in_rights or in_responsibilities:
                    # Skip very short lines (likely formatting artifacts)
                    if len(text) > 10:
                        if in_rights:
                            authorities_rights.append(text)
                        elif in_responsibilities:
                            authorities_responsibilities.append(text)

            elif current_subsection == 'functions':
                if text and text[0].isdigit() and ')' in text[:5]:
                    functions.append(text)

        elif current_chapter in [3, 4, 5]:
            additions.append(text)

    return {
        'general_provisions': '\n'.join(general_provisions),
        'tasks': tasks,
        'authorities_rights': authorities_rights,
        'authorities_responsibilities': authorities_responsibilities,
        'functions': functions,
        'additions': '\n'.join(additions)
    }


def parse_numbered_sections(doc):
    """Parse NUMBERED SECTIONS documents"""

    general_provisions = []
    tasks = []
    authorities_rights = []
    authorities_responsibilities = []
    functions = []
    additions = []

    current_section = None
    current_subsection = None
    in_rights = False
    in_responsibilities = False
    section_2_started = False
    section_3_or_higher = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if paragraph is bold (section headers are often bold, but not always)
        is_bold = any(run.bold for run in para.runs if run.bold)

        # Detect main sections (bold check relaxed - some files don't have bold headers)
        if text.startswith('1.') and 'Общие положения' in text:
            current_section = 1
            continue
        elif text.startswith('2.') and ('Задачи' in text or 'Миссия' in text):
            current_section = 2
            section_2_started = True
            current_subsection = None
            continue
        elif (len(text) < 100 and
              (text.startswith('3.') or text.startswith('4.') or
               text.startswith('5.') or text.startswith('6.')) and
              ('Статус' in text or 'Имущество' in text or 'Реорганизация' in text or
               'руководител' in text.lower())):
            current_section = 3
            section_3_or_higher = True
            additions.append(text)
            continue

        # Within section 2, detect subsections (numbered or unnumbered)
        if section_2_started and not section_3_or_higher:
            # Detect Задачи subsection (numbered like "14. Задачи:" or just "Задачи:")
            if ('Задачи' in text and len(text) < 100 and
                (text[0].isdigit() or text.startswith('Задачи'))):
                current_subsection = 'tasks'
                continue
            # Detect Полномочия subsection (or "Права и обязанности")
            elif (('Полномочия' in text or ('Права' in text and 'обязанности' in text)) and
                  len(text) < 150 and
                  (text[0].isdigit() or text.startswith('Полномочия') or text.startswith('Права'))):
                current_subsection = 'authorities'
                in_rights = False
                in_responsibilities = False
                # Check if "права:" is on the same line
                if 'права:' in text.lower() or 'права :' in text.lower():
                    in_rights = True
                continue
            # Detect Функции subsection
            elif ('Функции' in text and len(text) < 100 and
                  (text[0].isdigit() or text.startswith('Функции'))):
                current_subsection = 'functions'
                in_rights = False
                in_responsibilities = False
                continue

        # Extract based on current section
        if current_section == 1:
            general_provisions.append(text)

        elif current_section == 2 and section_2_started and not section_3_or_higher:
            if current_subsection == 'tasks':
                if text and text[0].isdigit() and ')' in text[:5]:
                    tasks.append(text)

            elif current_subsection == 'authorities':
                # Check for права/обязанности markers (may include numbering like "1) права:")
                if 'права:' in text.lower() or 'права :' in text.lower():
                    in_rights = True
                    in_responsibilities = False
                    continue
                elif 'обязанности:' in text.lower() or 'обязанности :' in text.lower():
                    in_rights = False
                    in_responsibilities = True
                    continue

                # In NUMBERED-SECTIONS, authorities are often NOT numbered
                # Collect all paragraphs under права or обязанности
                if in_rights or in_responsibilities:
                    # Skip very short lines (likely formatting artifacts)
                    if len(text) > 10:
                        if in_rights:
                            authorities_rights.append(text)
                        elif in_responsibilities:
                            authorities_responsibilities.append(text)

            elif current_subsection == 'functions':
                if text and text[0].isdigit() and ')' in text[:5]:
                    functions.append(text)
            elif current_subsection is None:
                # In section 2 but no subsection yet - might be combined section
                # Try to detect tasks/functions by pattern
                if text and text[0].isdigit() and ')' in text[:5]:
                    # Could be task or function - default to task for now
                    if not functions:  # If no functions yet, assume tasks
                        tasks.append(text)
                    else:
                        functions.append(text)

        elif section_3_or_higher:
            additions.append(text)

    return {
        'general_provisions': '\n'.join(general_provisions),
        'tasks': tasks,
        'authorities_rights': authorities_rights,
        'authorities_responsibilities': authorities_responsibilities,
        'functions': functions,
        'additions': '\n'.join(additions)
    }


def parse_custom(doc):
    """Parse CUSTOM documents using keyword-based extraction"""

    general_provisions = []
    tasks = []
    authorities_rights = []
    authorities_responsibilities = []
    functions = []
    additions = []

    current_zone = None
    in_rights = False
    in_responsibilities = False
    general_end = False
    tasks_section_started = False
    functions_section_started = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        text_lower = text.lower()

        # Detect zones by keywords
        if 'общие положения' in text_lower:
            current_zone = 'general'
            continue
        elif ('задачи' in text_lower and
              (len(text) < 100 or text.startswith('2.') or 'миссия' in text_lower)):
            current_zone = 'tasks'
            tasks_section_started = True
            general_end = True
            continue
        elif ('полномочия' in text_lower and len(text) < 100):
            current_zone = 'authorities'
            general_end = True
            in_rights = False
            in_responsibilities = False
            # Check if "права:" is on the same line
            if 'права:' in text_lower or 'права :' in text_lower:
                in_rights = True
            continue
        elif ('функции' in text_lower and (len(text) < 100 or text[0].isdigit())):
            current_zone = 'functions'
            functions_section_started = True
            general_end = True
            in_rights = False
            in_responsibilities = False
            continue
        elif ('имущество' in text_lower or
              'реорганизация' in text_lower or
              'статус' in text_lower) and len(text) < 100:
            current_zone = 'additions'
            additions.append(text)
            continue

        # Extract content
        if current_zone == 'general' and not general_end:
            general_provisions.append(text)

        elif current_zone == 'tasks':
            if text and text[0].isdigit() and ')' in text[:5]:
                tasks.append(text)

        elif current_zone == 'authorities':
            if 'права:' in text_lower or 'права :' in text_lower:
                in_rights = True
                in_responsibilities = False
                continue
            elif 'обязанности:' in text_lower or 'обязанности :' in text_lower:
                in_rights = False
                in_responsibilities = True
                continue

            # Collect both numbered (1., 2., 3.) and unnumbered items
            if in_rights or in_responsibilities:
                # Skip very short lines (likely formatting artifacts)
                if len(text) > 10:
                    if in_rights:
                        authorities_rights.append(text)
                    elif in_responsibilities:
                        authorities_responsibilities.append(text)

        elif current_zone == 'functions':
            if text and text[0].isdigit() and ')' in text[:5]:
                functions.append(text)

        elif current_zone == 'additions':
            additions.append(text)

    return {
        'general_provisions': '\n'.join(general_provisions),
        'tasks': tasks,
        'authorities_rights': authorities_rights,
        'authorities_responsibilities': authorities_responsibilities,
        'functions': functions,
        'additions': '\n'.join(additions)
    }


def parse_docx_universal(docx_path):
    """Universal parser that handles all document types"""

    doc = Document(docx_path)

    # Detect structure type
    structure_type = detect_structure_type(doc)

    print(f"\n📋 Detected structure: {structure_type}")

    # Parse based on structure
    if structure_type == 'CHAPTER-BASED':
        data = parse_chapter_based(doc)
    elif structure_type == 'NUMBERED-SECTIONS':
        data = parse_numbered_sections(doc)
    else:
        data = parse_custom(doc)

    return data


if __name__ == "__main__":
    # Test with a few different documents
    DATA_DIR = Path(__file__).parent.parent / "data"

    test_files = [
        "Положение УАиГ.docx",          # CHAPTER-BASED
        "Положение УВП 2025.docx",      # NUMBERED-SECTIONS
        "Положение УГА.docx",           # CUSTOM
    ]

    for filename in test_files:
        filepath = DATA_DIR / filename
        if filepath.exists():
            print(f"\n{'='*70}")
            print(f"TESTING: {filename}")
            print(f"{'='*70}")

            data = parse_docx_universal(filepath)

            print(f"\nResults:")
            print(f"  - General Provisions: {len(data['general_provisions'])} chars")
            print(f"  - Tasks: {len(data['tasks'])} items")
            print(f"  - Rights: {len(data['authorities_rights'])} items")
            print(f"  - Responsibilities: {len(data['authorities_responsibilities'])} items")
            print(f"  - Functions: {len(data['functions'])} items")
            print(f"  - Additions: {len(data['additions'])} chars")

            if data['tasks']:
                print(f"\n  First task: {data['tasks'][0][:80]}")
            if data['functions']:
                print(f"  First function: {data['functions'][0][:80]}")
