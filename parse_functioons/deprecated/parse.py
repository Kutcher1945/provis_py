"""
Parse "Положение УАиГ.docx" and convert to CSV with columns:
file, main_section, subsection, content
"""

import re
import csv
from docx import Document
from pathlib import Path

DOCX_PATH = Path(__file__).parent / "Положение УАиГ.docx"
OUTPUT_CSV = Path(__file__).parent / "parsed_output.csv"


def is_bold(paragraph):
    """Check if paragraph is bold (at least one bold run)"""
    return any(run.bold for run in paragraph.runs if run.bold)


def is_main_section(text, is_bold_para):
    """Identify main sections (Глава X, or bold headers)"""
    if not is_bold_para:
        return False
    # Check for "Глава X"
    if re.match(r'Глава\s+\d+\.?', text, re.IGNORECASE):
        return True
    # Check for other major bold headers
    if re.match(r'^[А-ЯЁ][а-яё\s]+$', text) and len(text) < 100:
        return True
    return False


def is_main_subsection(text):
    """Check if text is a main subsection (like '12. Задачи:' or '13. Полномочия:')"""
    # Match numbered items like "12. Задачи:" (double-digit or has colon/text after)
    match = re.match(r'^(\d+)\.\s+([А-ЯЁ])', text)
    if match:
        num = int(match.group(1))
        # If number is >= 10, it's likely a main subsection
        # Or if it has text after the number (not just another number)
        return num >= 10 or bool(re.match(r'^\d+\.\s+[А-ЯЁа-яё]+:', text))
    return False


def is_sub_item(text):
    """Check if text is a sub-item (like '1)', '2)' or sub-numbered '1.', '2.')"""
    # Match patterns like "1) " or "1. " (single digit with ) or .)
    match = re.match(r'^(\d+)[\.\)]\s+', text)
    if match:
        num = int(match.group(1))
        # Sub-items are typically single digits
        return num < 10 and (')')in text[:5] or not re.match(r'^\d+\.\s+[А-ЯЁ]', text)
    return False


def parse_document():
    """Parse the Word document and extract structured data"""
    doc = Document(DOCX_PATH)

    filename = DOCX_PATH.name
    current_main_section = ""
    current_subsection = ""
    records = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        bold = is_bold(para)

        # Check if this is a main section
        if is_main_section(text, bold):
            current_main_section = text
            current_subsection = ""
            # Add the main section itself as a record
            records.append({
                'file': filename,
                'main_section': current_main_section,
                'subsection': '',
                'content': ''
            })
            continue

        # Check if this is a main subsection (like "12. Задачи:")
        if is_main_subsection(text):
            current_subsection = text
            records.append({
                'file': filename,
                'main_section': current_main_section,
                'subsection': current_subsection,
                'content': ''
            })
            continue

        # Otherwise, it's content (including sub-items like "1)", "2)")
        records.append({
            'file': filename,
            'main_section': current_main_section,
            'subsection': current_subsection,
            'content': text
        })

    # Also extract tables
    for table_idx, table in enumerate(doc.tables, 1):
        table_text = f"[Таблица {table_idx}]\n"
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text += row_text + "\n"

        records.append({
            'file': filename,
            'main_section': current_main_section,
            'subsection': f'Таблица {table_idx}',
            'content': table_text.strip()
        })

    return records


def save_to_csv(records):
    """Save records to CSV with UTF-8 BOM for Excel compatibility"""
    with open(OUTPUT_CSV, 'w', newline='', encoding='utf-8-sig') as f:
        writer = csv.DictWriter(f, fieldnames=['file', 'main_section', 'subsection', 'content'])
        writer.writeheader()
        writer.writerows(records)

    print(f"Saved {len(records)} records to {OUTPUT_CSV}")


def main():
    print(f"Parsing {DOCX_PATH}...")
    records = parse_document()

    print(f"\nExtracted {len(records)} records")
    print(f"\nSample records:")
    for i, rec in enumerate(records[:10]):
        print(f"\n{i+1}. Main: {rec['main_section'][:50] if rec['main_section'] else '(none)'}")
        print(f"   Sub:  {rec['subsection'][:50] if rec['subsection'] else '(none)'}")
        print(f"   Text: {rec['content'][:80]}...")

    save_to_csv(records)
    print(f"\n✓ Done! CSV saved to {OUTPUT_CSV}")


if __name__ == '__main__':
    main()
